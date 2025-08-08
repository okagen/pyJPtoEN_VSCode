import os
from openai import OpenAI
from docx import Document
#from docx.oxml.text.paragraph import CT_P
#from docx.oxml.table import CT_Tc
import tiktoken
from tqdm import tqdm
from dotenv import load_dotenv
import configparser
import logging

#======================================================================
# 指定の翻訳モデルを使って、日本語テキストを英語に翻訳する
def translate_text(client, modelID, text, max_tokens=300, temperature=0.7):
    response = client.chat.completions.create(
        model=modelID,
        messages=[
            {"role":"system", "content":"You are a translator from Japanese to English."},
            {"role":"user",   "content":text}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    return response.choices[0].message.content.strip()

#======================================================================
# Word文書（Document）から、本文とすべての表内段落を順に取り出す
def iter_paragraphs(doc):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

#======================================================================
# 段落内の日本語テキストを、英語翻訳テキストに置き換える。
# 元の runs（書式）を可能な限り維持する。
def replace_text_preserve_styles(paragraph, new_text):
    full_text = paragraph.text
    if not full_text.strip():
        return
    if len(full_text) != len(new_text):
        # 長さが変わる時は runs に分解して置換
        # 簡略化のため、一度空runsクリアして全体追加
        for run in paragraph.runs:
            run.text = ""
        paragraph.add_run(new_text)
    else:
        # 長さが同じなら runsごとに直接置換
        idx = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run.text = new_text[idx:idx+run_len]
            idx += run_len

#======================================================================
# 指定のパスが既に存在する場合、自動で連番を追加して一意のファイル名を生成する
def uniquify(path):
    """
    指定された path が存在しなければそのまま返す。
    存在する場合、"filename (2).ext", "filename (3).ext", ... のようにして
    存在しない最初の名前を返す。
    """
    filename, extension = os.path.splitext(path)
    counter = 1
    new_path = path
    while os.path.exists(new_path):
        new_path = f"{filename} ({counter}){extension}"
        counter += 1
    return new_path

# Example usage
if __name__ == "__main__":
    # Configure the logger
    logging.basicConfig(
        filename='use.log',              # Output log file
        filemode='w',                    # Overwrite each time ('a' to append)
        level=logging.DEBUG,             # Log level (use INFO or WARNING for less verbosity)
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    logging.warning('Warning message')
    logging.error('Error message')

    load_dotenv()

    # Create a ConfigParser object
    config = configparser.ConfigParser()
    with open('config.ini', 'r', encoding='utf-8') as f:
        config.read_file(f)

    input_path =  config['input']['word_jp']
    output_path = f"output_{os.path.basename(input_path)}"
    output_path = uniquify(output_path)
    doc = Document(input_path)
    doc.save(output_path)
    out = Document(output_path)
    env_key = "OPENAI_API_KEY_005"
    #env_key = "OPENAI_API_KEY_SA_006"
    client = OpenAI(api_key=os.getenv(env_key))

    env_model_id = "OPENAI_API_MODEL_005"
    modelID = os.getenv(env_model_id)

    for para in tqdm(list(iter_paragraphs(out))):
        text = para.text.strip()
        if not text:
            continue
        try:
            print("===========================", env_key, " : ",env_model_id)
            print("【原文】" + text)

            translated = translate_text(client, modelID, text)

            print("-------------------------------------")
            print("【翻訳】" + translated)

            print("-------------------------------------")
            replace_text_preserve_styles(para, translated)
        except Exception as e:
            print("翻訳エラー:", e)

    out.save(output_path)
    print("完了:", output_path)

    #pyinstaller --onefile use.py
